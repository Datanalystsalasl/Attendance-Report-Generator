import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import calendar
import re
from datetime import datetime

# استخراج الـ Sheet ID من رابط Google Sheets
def extract_sheet_id(url):
    match = re.search(r'/d/([a-zA-Z0-9-_]+)', url)
    return match.group(1) if match else None

# دالة إنشاء التقرير
def generate_report(gsheetid, sheet_name="Total"):
    try:
        gsheet_url = f"https://docs.google.com/spreadsheets/d/{gsheetid}/gviz/tq?tqx=out:csv&sheet={sheet_name}"
        df = pd.read_csv(gsheet_url)

        # تنظيف الأعمدة
        df.columns = df.columns.str.strip()
        df = df.apply(lambda x: x.str.strip() if x.dtype == "object" else x)

        doc = Document()
        
        # إضافة صورة (لو متاحة)
        try:
            doc.add_picture('black.jpeg', width=Inches(1.25))
        except:
            pass
        
        # استخراج البيانات المطلوبة
        month = int(df.loc[0, 'month'])
        month_name = calendar.month_name[month]
        year = int(df.loc[0, 'Year'])
        total_working_days = df.loc[0, 'Total working days']

        def getLastdayOfmonth(month, year):
            last_day = calendar.monthrange(year, month)[1]
            return last_day

        getLastdayOfmonth(month, year)

        doc.add_heading(f'Attendance and Check-in Analysis - {month_name} {year}', 0)
        p = doc.add_paragraph('')
        p.add_run('Objective: ').bold = True
        p.add_run(f'To provide a summary of employee attendance, performance, and discipline for {month_name} {year} .')

        p = doc.add_paragraph('')
        p.add_run('Timeframe: ').bold = True
        p.add_run(f'from 1 {month_name} to {getLastdayOfmonth(month, year)} .')

        p = doc.add_paragraph('')
        p.add_run('Total number of working days: ').bold = True
        p.add_run(f'{total_working_days} days.')

        Overall_attendance = sum(df['Working Days']) - sum(df['NO.H']) - sum(df['Not Found days'])
        Overall_attendance_percentage = (Overall_attendance / sum(df['Working Days'])) * 100
        Overall_attendance_percentage = round(Overall_attendance_percentage, 2)

        p = doc.add_paragraph('')
        p.add_run('Overall attendance percentage: ').bold = True
        p.add_run(f'{Overall_attendance_percentage} %.')

        On_Site_percentage = round((sum(df['On site days']) / Overall_attendance) * 100, 2)
        on_site_per_day = round(sum(df['On site days']) / total_working_days)

        Remote_Days_percentage = round((sum(df['Remote Days']) / Overall_attendance) * 100, 2)
        Remote_Days_per_day = round(sum(df['Remote Days']) / total_working_days)

        p = doc.add_paragraph('')
        p.add_run('On Site VS Remotely: ').bold = True
        p.add_run(
            f'{On_Site_percentage} %. ({on_site_per_day} employee Per day) VS {Remote_Days_percentage} % ({Remote_Days_per_day} employee Per day ) ')

        a = pd.to_timedelta(df["Avg Check in"]).mean()

        hours, remainder = divmod(a.total_seconds(), 3600)
        minutes, seconds = divmod(remainder, 60)
        AvG_Check_In = f"{int(hours):02d}:{int(minutes):02d}:{int(seconds):02d}"

        a = pd.to_timedelta(df["AVG Check out"]).mean()

        hours, remainder = divmod(a.total_seconds(), 3600)
        minutes, seconds = divmod(remainder, 60)
        AvG_Check_Out = f"{int(hours):02d}:{int(minutes):02d}:{int(seconds):02d}"

        p = doc.add_paragraph('')
        p.add_run('Average check-in and check-out: ').bold = True
        p.add_run(f'{AvG_Check_In} VS {AvG_Check_Out} %.')

        p = doc.add_paragraph('')
        p.add_run('AVG Check in and Check out for every Department').bold = True

        avg_in_out_dept_sup = df.drop(['Id', 'Name', 'Working Days', 'Remote Days',
                                       'On site days', 'permission Days', 'NO.H', 'Not Found days',
                                       'Total working hours', 'Total Over time', 'total on time', 'total out time',
                                       'Total Delys hours',
                                       'month', 'Year', 'Total working days'], axis=1)

        avg_in_out_dept_sup['Avg Check in'] = pd.to_datetime(avg_in_out_dept_sup['Avg Check in'],
                                                             format='%H:%M:%S').dt.time
        avg_in_out_dept_sup['AVG Check out'] = pd.to_datetime(avg_in_out_dept_sup['AVG Check out'],
                                                              format='%H:%M:%S').dt.time

        avg_in_out_dept_sup_grouped = avg_in_out_dept_sup.groupby(['Dept', 'Sub Dept']).agg({
            'Avg Check in': lambda x: pd.to_timedelta(x.astype(str)).mean(),
            'AVG Check out': lambda x: pd.to_timedelta(x.astype(str)).mean()
        }).reset_index()
        avg_in_out_dept_sup_grouped['Avg Check in'] = avg_in_out_dept_sup_grouped['Avg Check in'].astype(
            'timedelta64[s]')
        avg_in_out_dept_sup_grouped['AVG Check out'] = avg_in_out_dept_sup_grouped['AVG Check out'].astype(
            'timedelta64[s]')

        avg_in_out_dept_sup_grouped['Avg Check in'] = avg_in_out_dept_sup_grouped['Avg Check in'].astype(str).str[-8:]
        avg_in_out_dept_sup_grouped['AVG Check out'] = avg_in_out_dept_sup_grouped['AVG Check out'].astype(str).str[-8:]

        table = doc.add_table(rows=1, cols=len(avg_in_out_dept_sup_grouped.columns))
        table.style = 'Table Grid'

        # إدراج عناوين الأعمدة مع التنسيق
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(avg_in_out_dept_sup_grouped.columns):
            hdr_cells[i].text = col_name

            # تنسيق النص ليكون بولد وأبيض
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)  # لون الخط أبيض

            # إضافة لون خلفية الهيدر (أسود)
            shading_elm = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
            hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)

        # إدراج بيانات الجدول
        for row in avg_in_out_dept_sup_grouped.values:
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row):
                row_cells[i].text = str(cell_value)

        p = doc.add_paragraph('')

        Total_Working_Hours = pd.to_timedelta(df['Total working hours']).sum()
        Total_Working_Hours = (Total_Working_Hours.days * 24) + (Total_Working_Hours.seconds / 3600)
        Total_Working_Hours = round(Total_Working_Hours, 2)

        p = doc.add_paragraph('')
        p.add_run('Total Working Hours: ').bold = True
        p.add_run(f'{Total_Working_Hours} Hour')

        Total_Over_Time = pd.to_timedelta(df['Total Over time']).sum()
        Total_Over_Time = (Total_Over_Time.days * 24) + (Total_Over_Time.seconds / 3600)
        Total_Over_Time = round(Total_Over_Time, 2)

        p = doc.add_paragraph('')
        p.add_run('Total Over Time: ').bold = True
        p.add_run(f'{Total_Over_Time} Hour')

        p = doc.add_paragraph('')
        p.add_run('Top 10 Employees by Performance Over Time').bold = True

        top_ten_emp_over_time = df.drop(['Id', 'Dept', 'Sub Dept', 'Working Days', 'Remote Days',
                                         'On site days', 'permission Days', 'NO.H', 'Not Found days',
                                         'Avg Check in', 'AVG Check out', 'total on time', 'total out time',
                                         'Total Delys hours', 'month', 'Year', 'Total working days'], axis=1)

        top_ten_emp_over_time['Total Over time'] = pd.to_timedelta(top_ten_emp_over_time['Total Over time'])

        top_ten_emp_over_time['Total Over time'] = round(
            top_ten_emp_over_time['Total Over time'].dt.total_seconds() / 3600, 2)
        final_top_ten_emp_over_time = top_ten_emp_over_time.sort_values('Total Over time', ascending=False).head(10)

        table = doc.add_table(rows=1, cols=len(final_top_ten_emp_over_time.columns))
        table.style = 'Table Grid'

        # إدراج عناوين الأعمدة مع التنسيق
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(final_top_ten_emp_over_time.columns):
            hdr_cells[i].text = col_name

            # تنسيق النص ليكون بولد وأبيض
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)  # لون الخط أبيض

            # إضافة لون خلفية الهيدر (أسود)
            shading_elm = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
            hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)

        # إدراج بيانات الجدول
        for row in final_top_ten_emp_over_time.values:
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row):
                row_cells[i].text = str(cell_value)

        p = doc.add_paragraph('')

        Total_number_of_permissions = df['permission Days'].sum()

        p = doc.add_paragraph('')
        p.add_run('Total number of permissions: ').bold = True
        p.add_run(f'{Total_number_of_permissions}')

        Total_number_of_Delaying = df['total out time'].sum()

        p = doc.add_paragraph('')
        p.add_run('Total number of Delaying: ').bold = True
        p.add_run(f'{Total_number_of_Delaying}')

        p = doc.add_paragraph('')
        p.add_run('Top 10 Employees with Delays by Hours ').bold = True

        tOP_ten_emp_delaying = df.drop(['Id', 'Dept', 'Sub Dept', 'Working Days', 'Remote Days',
                                        'On site days', 'permission Days', 'NO.H', 'Not Found days',
                                        'Total working hours', 'AVG working per day', 'Total Over time',
                                        'Avg Check in', 'AVG Check out', 'total on time',
                                        'month', 'Year', 'Total working days'], axis=1)

        tOP_ten_emp_delaying['Total Delys hours'] = pd.to_timedelta(tOP_ten_emp_delaying['Total Delys hours'])

        tOP_ten_emp_delaying['Total Delys hours'] = round(
            tOP_ten_emp_delaying['Total Delys hours'].dt.total_seconds() / 3600, 2)
        tOP_ten_emp_delaying = tOP_ten_emp_delaying.sort_values('Total Delys hours', ascending=False).head(10)

        # إضافة جدول
        table = doc.add_table(rows=1, cols=len(tOP_ten_emp_delaying.columns))
        table.style = 'Table Grid'

        # إدراج عناوين الأعمدة مع التنسيق
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(tOP_ten_emp_delaying.columns):
            hdr_cells[i].text = col_name

            # تنسيق النص ليكون بولد وأبيض
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)  # لون الخط أبيض

            # إضافة لون خلفية الهيدر (أسود)
            shading_elm = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
            hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)

        # إدراج بيانات الجدول
        for row in tOP_ten_emp_delaying.values:
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row):
                row_cells[i].text = str(cell_value)

        p = doc.add_paragraph('')

        p = doc.add_paragraph('')
        p.add_run('Top 15 Employees on time ').bold = True

        top_on_time = df.drop(['Id', 'Dept', 'Sub Dept', 'Working Days', 'Remote Days',
                               'On site days', 'permission Days', 'NO.H', 'Not Found days',
                               'Total working hours', 'AVG working per day', 'Total Over time',
                               'Avg Check in', 'AVG Check out', 'total out time',
                               'Total Delys hours', 'month', 'Year', 'Total working days'], axis=1)
        top_on_time = top_on_time.sort_values('total on time', ascending=False).head(15)

        # إضافة جدول
        table = doc.add_table(rows=1, cols=len(top_on_time.columns))
        table.style = 'Table Grid'

        # إدراج عناوين الأعمدة مع التنسيق
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(top_on_time.columns):
            hdr_cells[i].text = col_name

            # تنسيق النص ليكون بولد وأبيض
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)  # لون الخط أبيض

            # إضافة لون خلفية الهيدر (أسود)
            shading_elm = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
            hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)

        # إدراج بيانات الجدول
        for row in top_on_time.values:
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row):
                row_cells[i].text = str(cell_value)

        p = doc.add_paragraph('')

        p = doc.add_paragraph('')
        p.add_run('Number of employees have Not found days: ').bold = True
        filt = df['Not Found days'] != 0
        Number = df.loc[filt].shape[0]
        p.add_run(f'{Number}')

        filt2 = (df['Not Found days'] < 0) | (df['Not Found days'] > 9)
        Fingerprint_Authentication_Issues = df.loc[filt2, ['Name', 'Not Found days']]
        Fingerprint_Authentication_Issues = Fingerprint_Authentication_Issues.sort_values('Not Found days',
                                                                                          ascending=False)

        p = doc.add_paragraph('')
        p.add_run('Employees: Fingerprint Authentication Issues').bold = True

        # إضافة جدول
        table = doc.add_table(rows=1, cols=len(Fingerprint_Authentication_Issues.columns))
        table.style = 'Table Grid'

        # إدراج عناوين الأعمدة مع التنسيق
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(Fingerprint_Authentication_Issues.columns):
            hdr_cells[i].text = col_name

            # تنسيق النص ليكون بولد وأبيض
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)  # لون الخط أبيض

            # إضافة لون خلفية الهيدر (أسود)
            shading_elm = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
            hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)

        # إدراج بيانات الجدول
        for row in Fingerprint_Authentication_Issues.values:
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row):
                row_cells[i].text = str(cell_value)

        p = doc.add_paragraph('')
        file_name = f"Attendance_Report_{month_name}_{year}.docx"
        doc.save(file_name)
        return file_name

    except Exception as e:
        st.error(f"❌ حدث خطأ أثناء إنشاء التقرير: {e}")
        return None

# --- واجهة Streamlit ---
st.image("black.jpeg", width=200)
st.title("📊 Attendance Report Generator")
st.write("Make sure sheet linke accessible 'Anyone with the link'  ")

# إدخال رابط جوجل شيت
sheet_link = st.text_input("put sheet link:", placeholder="https://docs.google.com/spreadsheets/d/your-sheet-id")

if sheet_link:
    gsheetid = extract_sheet_id(sheet_link)

    if gsheetid:
        st.success("done")

        if st.button("Create Report"):
            report_file = generate_report(gsheetid)
            if report_file:
                with open(report_file, "rb") as file:
                    st.download_button("Dwonload Reportر", file, file_name=report_file)
    else:
        st.error("❌ الرابط غير صحيح، تأكد من إدخال رابط Google Sheet بشكل صحيح.")
